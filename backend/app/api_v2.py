"""
Enhanced FastAPI routes for production Document AI Assistant.
Features: Auto-summaries, multi-document support, chat history, and more.
"""

from fastapi import APIRouter, UploadFile, File, Depends, HTTPException, Header
from fastapi.responses import JSONResponse
from sqlalchemy.orm import Session
from sqlalchemy import desc
from google.api_core import exceptions as google_exceptions
from pydantic import BaseModel
from typing import Optional, List
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document
import os
import re
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from dotenv import load_dotenv

from . import embeddings
from .database import get_db
from .models import Document as DocumentModel, ChatMessage
from .prompts import get_chat_prompt

load_dotenv()

router = APIRouter()

UPLOAD_FOLDER = "./data/uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# --------------------------
# Pydantic Models
# --------------------------
class QuestionRequest(BaseModel):
    question: str
    document_id: Optional[str] = None  # If None, search all documents (deprecated, use document_ids)
    document_ids: Optional[List[str]] = None  # New: list of document IDs for multi-document search
    use_chat_history: bool = True


class UploadResponse(BaseModel):
    document_id: str
    filename: str
    summary: str
    chunk_count: int


class DocumentListItem(BaseModel):
    id: str
    filename: str
    upload_time: str
    summary: str
    chunk_count: int
    is_active: bool


class BulkDeleteRequest(BaseModel):
    document_ids: List[str]


# --------------------------
# Health Check
# --------------------------
@router.get("/health")
async def health():
    return {"status": "ok", "version": "2.0"}


# --------------------------
# Text Processing Functions
# --------------------------
def preprocess_text(text: str) -> str:
    """Normalize text: lowercase, remove extra whitespace, remove punctuation."""
    text = text.lower()
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^\w\s]", "", text)
    return text


def remove_stopwords(text: str) -> str:
    """Remove common English stopwords."""
    stop_words = set(stopwords.words("english"))
    word_tokens = word_tokenize(text)
    filtered_text = [w for w in word_tokens if not w.lower() in stop_words]
    return " ".join(filtered_text)


def chunk_text(text: str, chunk_size=500, chunk_overlap=150):
    """Split text into overlapping chunks for embedding."""
    words = text.split()
    if not words:
        return []
    return [
        " ".join(words[i : i + chunk_size])
        for i in range(0, len(words), chunk_size - chunk_overlap)
    ]


def extract_text_from_file(file_path: str, filename: str) -> str:
    """Extract text from PDF, DOCX, or TXT files."""
    text = ""
    try:
        if filename.endswith(".pdf"):
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        elif filename.endswith(".docx"):
            full_text = []
            doc = Document(file_path)
            for para in doc.paragraphs:
                full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            full_text.append(paragraph.text)
            text = "\n".join(full_text)
        elif filename.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        else:
            raise ValueError(f"Unsupported file type: {filename}")
    except Exception as e:
        print(f"Error extracting text from {filename}: {e}")
        raise
    
    return text


# --------------------------
# Generate Summary using Gemini
# --------------------------
def generate_summary(text: str, filename: str) -> str:
    """
    Generate a concise, professional summary using Gemini API.
    Designed to be quick and informative for the user.
    """
    try:
        # Limit text for summary generation to avoid token limits
        text_preview = text[:5000]  # First 5000 chars for summary context
        
        prompt = f"""
        Please provide a concise professional summary (2-3 sentences) of the following document titled "{filename}".
        The summary should be informative and capture the main topic/purpose.
        
        Document excerpt:
        {text_preview}
        
        Summary:
        """
        
        model = genai.GenerativeModel("gemini-2.5-flash")
        response = model.generate_content(prompt)
        return response.text.strip()
    
    except Exception as e:
        print(f"Error generating summary: {e}")
        return f"Could not generate summary for {filename}"


# --------------------------
# File Upload with Summary
# --------------------------

@router.post("/upload", response_model=UploadResponse)
async def upload_file(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    x_device_id: Optional[str] = Header(None, convert_underscores=False)
):
    """
    Upload and process a document synchronously.
    The request will complete only after the document is fully processed.
    """
    print("--- Starting file upload ---")
    print(f"Received X-Device-Id on upload: {x_device_id}")
    file_location = os.path.join(UPLOAD_FOLDER, file.filename)
    print(f"File location: {file_location}")
    
    # Create initial document record
    doc = DocumentModel(
        filename=file.filename,
        file_path=file_location,
        summary="Processing...",
        is_active=True,
        device_id=x_device_id,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    print(f"Created document record with ID: {doc.id}")

    try:
        # Save file temporarily
        print("Saving uploaded file...")
        with open(file_location, "wb") as f:
            f.write(await file.read())
        print("File saved.")

        # --- Synchronous Processing ---
        
        # 1. Extract text
        print("Extracting text from file...")
        text = extract_text_from_file(file_location, file.filename)
        if not text.strip():
            raise ValueError("No text could be extracted from the file.")
        print(f"Text extracted. Length: {len(text)} characters.")
            
        # 2. Generate summary
        print("Generating summary...")
        summary = generate_summary(text, file.filename)
        print("Summary generated.")
        
        # 3. Preprocess and chunk text
        print("Preprocessing and chunking text...")
        processed_text = preprocess_text(text)
        text_without_stopwords = remove_stopwords(processed_text)
        chunks = chunk_text(text_without_stopwords, chunk_size=500, chunk_overlap=150)
        
        if not chunks:
            raise ValueError("Text was extracted, but no processable chunks were generated.")
        print(f"Text chunked into {len(chunks)} chunks.")

        # 4. Add chunks to the vector index
        print("Adding chunks to vector index...")
        embeddings.add_chunks_to_index(doc.id, chunks)
        print("Chunks added to index.")

        # 5. Finalize database update
        print("Finalizing database update...")
        doc.summary = summary
        doc.chunk_count = len(chunks)
        doc.document_size = len(text)
        # Ensure device_id stored for visibility/debug
        if x_device_id:
            doc.device_id = x_device_id
        db.commit()
        db.refresh(doc)
        
        print(f"✓ Document {file.filename} processed and indexed successfully.")

        return UploadResponse(
            document_id=doc.id,
            filename=doc.filename,
            summary=doc.summary,
            chunk_count=doc.chunk_count
        )

    except Exception as e:
        # If anything fails, roll back the initial document creation
        db.delete(doc)
        db.commit()
        print(f"✗ Error processing {file.filename}: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        # Clean up the temporarily saved file
        print("Cleaning up temporary file...")
        if os.path.exists(file_location):
            os.remove(file_location)
        print("--- File upload finished ---")


# --------------------------
# Document Management Endpoints
# --------------------------
@router.get("/documents", response_model=List[DocumentListItem])
async def list_documents(
    db: Session = Depends(get_db),
    x_device_id: Optional[str] = Header(None, convert_underscores=False),
):
    """Get list of uploaded documents. If `X-Device-Id` header is provided, filter by device."""
    print(f"List documents called. X-Device-Id: {x_device_id}")
    query = db.query(DocumentModel)
    if x_device_id:
        query = query.filter(DocumentModel.device_id == x_device_id)
    docs = query.order_by(desc(DocumentModel.upload_time)).all()
    return [DocumentListItem(**doc.to_dict()) for doc in docs]


@router.get("/documents/{document_id}")
async def get_document(document_id: str, db: Session = Depends(get_db)):
    """Get details of a specific document."""
    doc = db.query(DocumentModel).filter(DocumentModel.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc.to_dict()


@router.post("/documents/{document_id}/set-active")
async def set_active_document(document_id: str, db: Session = Depends(get_db)):
    """Mark a document as the active context for questions."""
    # Deactivate all documents
    db.query(DocumentModel).update({"is_active": False})
    
    # Activate the selected document
    doc = db.query(DocumentModel).filter(DocumentModel.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    doc.is_active = True
    db.commit()
    
    return {"message": f"Document '{doc.filename}' is now active"}


@router.delete("/documents/{document_id}")
async def delete_document(document_id: str, db: Session = Depends(get_db)):
    """Delete a document and its embeddings."""
    doc = db.query(DocumentModel).filter(DocumentModel.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Delete from FAISS
    embeddings.delete_index(document_id)
    
    # Delete from database
    db.delete(doc)
    db.commit()
    
    return {"message": f"Document '{doc.filename}' deleted successfully"}


@router.post("/documents/bulk-delete")
async def bulk_delete_documents(req: BulkDeleteRequest, db: Session = Depends(get_db)):
    """Delete multiple documents and their embeddings in one request."""
    deleted = 0
    try:
        for document_id in req.document_ids:
            doc = db.query(DocumentModel).filter(DocumentModel.id == document_id).first()
            if not doc:
                continue
            # delete faiss index
            embeddings.delete_index(document_id)
            db.delete(doc)
            deleted += 1
        db.commit()
        return {"deleted": deleted}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/debug/documents")
async def debug_list_all_documents(db: Session = Depends(get_db)):
    """Debug endpoint: return all documents with their stored device_id (no filtering).
    Intended for debugging deployments only. Remove or protect in production."""
    docs = db.query(DocumentModel).order_by(desc(DocumentModel.upload_time)).all()
    return [doc.to_dict() for doc in docs]


# --------------------------
# Chat History Endpoints
# --------------------------
@router.get("/chat-history/{document_id}")
async def get_chat_history(document_id: str, db: Session = Depends(get_db)):
    """Get chat history for a document."""
    messages = db.query(ChatMessage).filter(
        ChatMessage.document_id == document_id
    ).order_by(ChatMessage.timestamp).all()
    
    return [msg.to_dict() for msg in messages]


@router.delete("/chat-history/{document_id}")
async def clear_chat_history(document_id: str, db: Session = Depends(get_db)):
    """Clear chat history for a document."""
    db.query(ChatMessage).filter(ChatMessage.document_id == document_id).delete()
    db.commit()
    return {"message": "Chat history cleared"}


# --------------------------
# Ask Question with Chat Memory
# --------------------------
@router.post("/ask")
async def ask_question(
    q: QuestionRequest,
    db: Session = Depends(get_db)
):
    """
    Ask a question about the document(s).
    Supports conversation history and multi-document search.
    Can search in:
    - Single document (document_id provided)
    - Multiple documents (document_ids list provided)
    - All documents (neither provided)
    """
    try:
        # Determine which document(s) to search
        search_doc_ids = None
        
        if q.document_ids:
            # Multi-document search
            search_doc_ids = q.document_ids
            # Validate all documents exist
            for doc_id in search_doc_ids:
                if not db.query(DocumentModel).filter(DocumentModel.id == doc_id).first():
                    raise HTTPException(status_code=404, detail=f"Document {doc_id} not found")
        elif q.document_id:
            # Single document search (backward compatible)
            if not db.query(DocumentModel).filter(DocumentModel.id == q.document_id).first():
                raise HTTPException(status_code=404, detail="Document not found")
            search_doc_ids = [q.document_id]
        
        # Search for relevant chunks from selected document(s)
        relevant_chunks = []
        if search_doc_ids:
            # Search each document and collect chunks
            for doc_id in search_doc_ids:
                chunks = embeddings.search(q.question, document_id=doc_id, top_k=2)
                relevant_chunks.extend(chunks)
        else:
            # Search all documents (global search)
            relevant_chunks = embeddings.search(q.question, document_id=None, top_k=5)
        
        if not relevant_chunks:
            return {
                "answer": "I don't have any relevant information to answer that question. Please upload a document first.",
                "source_chunks": []
            }
        
        # Format context with labels so the LLM can reference them
        context_text = ""
        for i, chunk in enumerate(relevant_chunks, 1):
            context_text += f"[Excerpt {i}]:\n{chunk}\n\n"
        
        # Build conversation history for context
        conversation_context = ""
        if q.use_chat_history and search_doc_ids:
            # Get recent messages from first document in the list
            first_doc_id = search_doc_ids[0]
            recent_messages = db.query(ChatMessage).filter(
                ChatMessage.document_id == first_doc_id
            ).order_by(desc(ChatMessage.timestamp)).limit(6).all()
            
            recent_messages.reverse()
            for msg in recent_messages:
                role = "User" if msg.role == "user" else "Assistant"
                conversation_context += f"{role}: {msg.content}\n"
        
        # Build prompt with context
        prompt = get_chat_prompt(q.question, context_text, conversation_context)
        
        # Generate response using Gemini
        model = genai.GenerativeModel("gemini-2.5-flash")
        response = model.generate_content(prompt)
        answer = response.text
        
        # Save to chat history if single document provided
        if search_doc_ids and len(search_doc_ids) == 1:
            doc_id = search_doc_ids[0]
            user_msg = ChatMessage(
                document_id=doc_id,
                role="user",
                content=q.question
            )
            assistant_msg = ChatMessage(
                document_id=doc_id,
                role="assistant",
                content=answer
            )
            db.add(user_msg)
            db.add(assistant_msg)
            db.commit()
        
        return {
            "answer": answer,
            "source_chunks": relevant_chunks
        }
    
    except google_exceptions.ServiceUnavailable as e:
        error_message = "Could not connect to Google's AI service. Please try again."
        return JSONResponse(status_code=503, content={"answer": error_message})
    
    except google_exceptions.RetryError as e:
        error_message = "Request to AI service timed out. Please try again."
        return JSONResponse(status_code=504, content={"answer": error_message})
    
    except Exception as e:
        print(f"Error in ask_question: {e}")
        return JSONResponse(status_code=500, content={"error": str(e)})


# --------------------------
# Summary Endpoint
# --------------------------
@router.get("/documents/{document_id}/summary")
async def get_document_summary(document_id: str, db: Session = Depends(get_db)):
    """Get or regenerate summary for a specific document."""
    try:
        doc = db.query(DocumentModel).filter(DocumentModel.id == document_id).first()
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        
        return {
            "document_id": doc.id,
            "filename": doc.filename,
            "summary": doc.summary,
            "chunk_count": doc.chunk_count
        }
    
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})


@router.post("/documents/{document_id}/summary/regenerate")
async def regenerate_document_summary(document_id: str, db: Session = Depends(get_db)):
    """Regenerate summary and embeddings for a document synchronously."""
    doc = db.query(DocumentModel).filter(DocumentModel.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    file_location = doc.file_path
    if not os.path.exists(file_location):
        raise HTTPException(status_code=404, detail=f"File not found at path: {file_location}. Cannot regenerate.")

    try:
        # Mark as processing
        doc.summary = "Processing..."
        db.commit()

        # --- Synchronous Processing ---
        
        # 1. Extract text
        text = extract_text_from_file(file_location, doc.filename)
        if not text.strip():
            raise ValueError("No text could be extracted from the file for regeneration.")
            
        # 2. Generate summary
        summary = generate_summary(text, doc.filename)
        
        # 3. Preprocess and chunk text
        processed_text = preprocess_text(text)
        text_without_stopwords = remove_stopwords(processed_text)
        chunks = chunk_text(text_without_stopwords, chunk_size=500, chunk_overlap=150)
        
        if not chunks:
            raise ValueError("Text was extracted, but no processable chunks were generated for regeneration.")

        # 4. Delete old index and add new chunks
        embeddings.delete_index(doc.id)
        embeddings.add_chunks_to_index(doc.id, chunks)

        # 5. Finalize database update
        doc.summary = summary
        doc.chunk_count = len(chunks)
        doc.document_size = len(text)
        db.commit()
        
        print(f"✓ Document {doc.filename} regenerated and indexed successfully.")
        
        return {"message": "Regeneration successful", "document_id": doc.id}

    except Exception as e:
        # Revert summary on failure
        doc.summary = "Regeneration failed. Please try again."
        db.commit()
        print(f"✗ Error regenerating {doc.filename}: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    
